import os, re
from pypdf import PdfReader
src = r'C:\Users\SIDDHE~1\AppData\Local\Temp\claude\F--Projects-Siddhesh-Thapa\e5be9644-09e9-4563-861f-7d2d398064ec\scratchpad\source'
out = r'C:\Users\SIDDHE~1\AppData\Local\Temp\claude\F--Projects-Siddhesh-Thapa\e5be9644-09e9-4563-861f-7d2d398064ec\scratchpad\extracted'
os.makedirs(out, exist_ok=True)

def save(name, text):
    text = re.sub(r'\n{3,}', '\n\n', text).strip()
    p = os.path.join(out, name + '.txt')
    open(p, 'w', encoding='utf-8').write(text)
    print('%-46s %7d chars' % (name, len(text)))

for f in sorted(os.listdir(src)):
    p = os.path.join(src, f)
    base = os.path.splitext(f)[0]
    ext = os.path.splitext(f)[1].lower()
    try:
        if ext == '.pdf':
            r = PdfReader(p)
            save(base, '\n'.join((pg.extract_text() or '') for pg in r.pages))
        elif ext == '.docx':
            import docx
            d = docx.Document(p)
            parts = [para.text for para in d.paragraphs]
            for t in d.tables:
                for row in t.rows:
                    parts.append(' | '.join(c.text for c in row.cells))
            save(base, '\n'.join(parts))
        elif ext == '.rtf':
            from striprtf.striprtf import rtf_to_text
            save(base, rtf_to_text(open(p, encoding='utf-8', errors='ignore').read()))
    except Exception as e:
        print('%-46s FAILED %s' % (base, e))